import re
import pandas as pd
from docx import Document

def parse_nda(doc_path):
    """
    Parse NDA .docx/.doc file and extract sentences.
    Returns DataFrame with original and normalized sentences.
    """

    # try loading doc
    try:
        doc = Document(doc_path)
    except Exception as e:
        print("Error loading file:", e)
        return None

    # END patterns to detect body end
    END_PATTERNS = [
        r'in\s+witness\s+whereof',
        r'in\s+witness\s+thereof',
        r'the\s+signatures\s+follow',
        r'signature\s+page\s+follows',
        r'signed\s+on\s+behalf\s+of',
        r'signed\s+for\s+and\s+on\s+behalf\s+of',
        r'each\s+acting\s+under\s+due\s+and\s+proper\s+authority'
    ]
    end_pattern = re.compile("|".join(END_PATTERNS), re.IGNORECASE)

    paragraphs = []
    found_end = False

    # 1) search in paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        paragraphs.append(text)

        if end_pattern.search(text):
            found_end = True
            break

    # 2) fallback to tables 
    if not found_end:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    txt = cell.text.strip()
                    if not txt:
                        continue
                    paragraphs.append(txt)
                    if end_pattern.search(txt):
                        found_end = True
                        break
                if found_end:
                    break
            if found_end:
                break

    if not paragraphs:
        return None

    # ==========================
    # sentence split
    # ==========================
    raw_sents = []
    for para_text in paragraphs:
        split_sents = re.split(r"(?<=[.?!])\s+(?=[A-Z(])", para_text)
        raw_sents.extend([s.strip() for s in split_sents if s.strip()])

    parsed = []
    main_clause = None

    for sent in raw_sents:
        if len(sent.split()) < 5 and sent.endswith("."):
            continue

        elif sent.endswith(":"):
            main_clause = sent

        elif main_clause:
            parsed.append(f"{main_clause} {sent}")
            if sent.endswith("."):
                main_clause = None

        else:
            parsed.append(sent)

    # ==========================
    # normalize
    # ==========================
    def normalize(x):
        x = x.lower()
        x = re.sub(r"\s+", " ", x)
        x = re.sub(r"^[a-z]\)|^\([a-z]\)|^\d+[\.\)]", "", x)
        return x.strip()

    normalized = [normalize(s) for s in parsed]

    df = pd.DataFrame({
        "original_sentence": parsed,
        "normalized_sentence": normalized
    })

    return df
